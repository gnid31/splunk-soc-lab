#!/usr/bin/env python3
"""
DFIR Report Generator — SOC Lab
Generates a professional Word document from attack simulation evidence.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── helpers ──────────────────────────────────────────────────────────────────
def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = RGBColor(*color)
    return h

def add_para(doc, text, bold=False, italic=False, size=11, color=None, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_code(doc, text, title=None):
    """Monospace code block with grey background."""
    if title:
        tp = doc.add_paragraph()
        tr = tp.add_run(f"  {title}")
        tr.bold = True
        tr.font.size = Pt(9)
        tr.font.color.rgb = RGBColor(0x44,0x44,0x44)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)

    # shade background
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)

    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1E,0x1E,0x1E)
    return p

def add_table(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # header row
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        # blue background
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1F3A6E')
        tc_pr.append(shd)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

    for row_data in rows:
        row = tbl.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(9)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Cm(w)
    return tbl

def spacer(doc, n=1):
    for _ in range(n):
        doc.add_paragraph()

# ── main document ─────────────────────────────────────────────────────────────
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ═══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════════════════════
spacer(doc, 4)
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("DFIR INCIDENT INVESTIGATION REPORT")
r.bold = True; r.font.size = Pt(24)
r.font.color.rgb = RGBColor(0x1F,0x3A,0x6E)

spacer(doc)
t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t2.add_run("SOC Lab — Windows 10 Endpoint Attack Simulation")
r2.font.size = Pt(16); r2.bold = True
r2.font.color.rgb = RGBColor(0x44,0x44,0x44)

spacer(doc, 2)
meta = [
    ("Incident ID",      "INC-2026-0817-001"),
    ("Date/Time",        "2026-08-17  21:44:05 – 21:45:28 ICT (14:44:05 – 14:45:28 UTC)"),
    ("Analyst",          "SOC Lab Training"),
    ("Host",             "DESKTOP-L7FCMBQ  (192.168.154.164, Windows 10 19045)"),
    ("User Account",     "gnid  (Local Administrator)"),
    ("SIEM",             "Elastic Stack 8.14.3 / ELK on VPS 43.228.215.234"),
    ("Log Source",       "Sysmon 15.21 (SwiftOnSecurity config) + Winlogbeat 8.19.0"),
    ("Classification",   "CONFIDENTIAL – Training Purpose Only"),
]
for k, v in meta:
    row = doc.add_paragraph()
    row.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rb = row.add_run(f"{k}: "); rb.bold = True; rb.font.size = Pt(11)
    rv = row.add_run(v);       rv.font.size = Pt(11)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "1. Executive Summary", 1, (0x1F,0x3A,0x6E))
add_para(doc,
    "On 2026-08-17 at 21:44 ICT, an attack simulation was executed on Windows 10 endpoint "
    "DESKTOP-L7FCMBQ. The attack replicated a realistic, multi-stage intrusion lifecycle: "
    "initial access via encoded PowerShell, post-exploitation discovery, persistence establishment, "
    "credential access (LSASS handle), defense evasion (log clear), and a C2 beacon attempt. "
    "All 7 attack stages were successfully detected and traced through Sysmon event logs forwarded "
    "to the ELK stack via Winlogbeat.")

spacer(doc)
add_heading(doc, "Findings Summary", 2)
summary_rows = [
    ("1","Execution",        "T1059.001","PowerShell –EncodedCommand executed","CONFIRMED","EID-4104, EID-1"),
    ("2","Discovery",        "T1087, T1082, T1016","whoami, systeminfo, ipconfig, netstat, tasklist","CONFIRMED","EID-1"),
    ("3","Persistence",      "T1547.001","Registry Run key: WindowsUpdateSvc → calc.exe","CONFIRMED","EID-13"),
    ("3b","Persistence",     "T1053.005","Scheduled Task: MicrosoftEdgeUpdateTaskUA","CONFIRMED","EID-4103"),
    ("4","Defense Evasion",  "T1070.001","wevtutil cl Application (log cleared)","CONFIRMED","EID-4104"),
    ("4b","Defense Evasion", "T1036",    "Masquerade file: svchost32.exe.log in %TEMP%","CONFIRMED","EID-11"),
    ("5","Credential Access","T1003.001","LSASS handle open (mask=0x1010) by powershell.exe","CONFIRMED","EID-10"),
    ("6","Lateral Movement", "T1021.002","SMB admin share probe: net use \\127.0.0.1\\C$","EVIDENCE","EID-1"),
    ("7","C&C",              "T1071.001","HTTP beacon to 93.184.216.34:80 (×5)","EVIDENCE","EID-3, DNS"),
]
add_table(doc,
    ["#","Tactic","Technique","Activity","Status","Evidence"],
    summary_rows,
    [0.8, 2.5, 2.2, 5.5, 2.2, 3.5])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 2. ANALYST THOUGHT PROCESS
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "2. Analyst Thought Process & Investigation Approach", 1, (0x1F,0x3A,0x6E))
add_para(doc,
    "Phần này mô tả tư duy và phương pháp tiếp cận của analyst khi điều tra incident — "
    "đây là phần quan trọng nhất để học DFIR, không chỉ biết chạy query mà còn phải "
    "hiểu TẠI SAO chạy query đó và nhìn vào ĐÂU tiếp theo.")

spacer(doc)
add_heading(doc, "2.1  Điểm khởi đầu — Alert hoặc anomaly", 2)
add_para(doc,
    "Một analyst thực tế sẽ nhận được cảnh báo từ detection rule (ví dụ: Splunk/ELK alert "
    "'PowerShell EncodedCommand'). Câu hỏi đầu tiên luôn là:")
for q in [
    "1. Alert này có THẬT không hay false positive?",
    "2. Nếu thật, đây là giai đoạn nào của attack lifecycle (Initial Access? Execution? C2?)?",
    "3. SCOPE: bao nhiêu host bị ảnh hưởng? Có lateral movement không?",
    "4. TIMELINE: khi nào bắt đầu? Attacker còn đang active không?",
]:
    add_para(doc, q, indent=0.5)

spacer(doc)
add_heading(doc, "2.2  Framework tư duy: Diamond Model + MITRE ATT&CK", 2)
add_para(doc,
    "Analyst không trace log ngẫu nhiên — dùng framework để có hướng đi:")
add_para(doc,
    "• MITRE ATT&CK: mỗi alert map tới 1 technique → biết attacker đang ở tactic nào → "
    "dự đoán bước tiếp theo → chủ động hunt thay vì reactive.", indent=0.5)
add_para(doc,
    "• Kill Chain: nếu thấy Execution → phải ngay lập tức hunt Persistence và "
    "Credential Access vì đó là bước attacker làm tiếp theo.", indent=0.5)
add_para(doc,
    "• Ví dụ tư duy trong incident này:", indent=0.5)
add_code(doc,
    "Thấy EID-4104 (PS ScriptBlock) → suspicion: attack script running\n"
    "→ Pivot sang EID-1 (ProcessCreate): powershell.exe sinh ra cmd.exe\n"
    "→ cmd.exe chạy whoami, systeminfo → Discovery phase confirmed\n"
    "→ Hunt persistence: EID-13 (Registry) → tìm thấy Run key\n"
    "→ Hunt cred access: EID-10 (ProcessAccess lsass) → 0x1010 mask confirmed\n"
    "→ Scope check: EID-3 (network) → C2 beacon pattern")

spacer(doc)
add_heading(doc, "2.3  Tips & Tricks khi trace log trên ELK", 2)

tips = [
    ("TIP 1: Luôn bắt đầu từ time range chính xác",
     "Đừng search 'All time' — tốn CPU và nhiễu. Xác định attack window trước:\n"
     "• Lấy timestamp từ alert đầu tiên\n"
     "• Mở rộng ±15 phút để bắt pre/post activity\n"
     "• KQL: @timestamp >= \"2026-08-17T14:44:00Z\" and @timestamp <= \"2026-08-17T15:00:00Z\""),

    ("TIP 2: Pivot theo ProcessGuid, không theo tên process",
     "Tên process dễ giả mạo (T1036). ProcessGuid là unique per-process-lifetime:\n"
     "• Tìm ProcessGuid của process đáng ngờ từ EID-1\n"
     "• Pivot: winlog.event_data.SourceProcessGUID: \"{guid}\" để tìm mọi activity từ process đó\n"
     "• Pivot: winlog.event_data.ParentProcessGuid: \"{guid}\" để tìm children"),

    ("TIP 3: Đọc Parent→Child process tree",
     "Attack chain hiện rõ trong process tree:\n"
     "• sshd.exe → cmd.exe → powershell.exe → [attack script]\n"
     "• Legitimate: WINWORD.EXE không nên spawn powershell.exe\n"
     "• Red flag: powershell.exe spawn cmd.exe spawn whoami.exe (discovery)"),

    ("TIP 4: Dùng message field khi mapping chưa có keyword",
     "Old Winlogbeat (8.x) đôi khi map event_id là text, không aggregate được.\n"
     "Giải pháp: dùng event.code hoặc match trên message field:\n"
     "• event.code: \"1\" AND message: \"EncodedCommand\"\n"
     "• message: \"TargetObject\" AND message: \"CurrentVersion\\\\Run\""),

    ("TIP 5: Biết các access mask quan trọng (EID-10)",
     "Không phải mọi LSASS access đều là Mimikatz. Phân biệt:\n"
     "• 0x1010 = PROCESS_QUERY_LIMITED_INFO | VM_READ → Mimikatz-style\n"
     "• 0x1400 = PROCESS_QUERY_INFO | VM_READ → WMI/System thường\n"
     "• 0x0040 = PROCESS_DUP_HANDLE → injection\n"
     "• Context: ai access? svchost, MsMpEng → benign; powershell → alert"),

    ("TIP 6: Hunt by keyword trước, filter sau",
     "Workflow hiệu quả:\n"
     "• Bước 1: search rộng → message: \"lsass\" → thấy 548 events\n"
     "• Bước 2: filter noise → NOT message: \"wmiprvse\" NOT message: \"MsMpEng\"\n"
     "• Bước 3: zoom in → message: \"0x1010\" → thấy 2 events, 1 là powershell.exe\n"
     "• Sai lầm phổ biến: filter quá sớm → miss evidence"),

    ("TIP 7: Luôn verify context của process lạ",
     "Khi thấy process lạ, verify theo checklist:\n"
     "• Đường dẫn có đúng không? (svchost.exe chỉ run từ C:\\Windows\\System32)\n"
     "• Parent process có hợp lệ không?\n"
     "• Command line argument có suspicious không?\n"
     "• Có network connection sau khi start không? (C2 indicator)\n"
     "• Có file create/registry write không? (persistence indicator)"),

    ("TIP 8: Dùng Discover 'Surrounding documents'",
     "Trong Kibana Discover, click vào 1 event → 'View surrounding documents'\n"
     "→ Xem ±50 events trước/sau theo thời gian → rất hữu ích để trace sequence\n"
     "Equivalent query:\n"
     "host.name: \"DESKTOP-L7FCMBQ\" AND @timestamp: [T-30s TO T+30s]"),
]

for title, content in tips:
    add_para(doc, f"▶  {title}", bold=True, color=(0x1F,0x3A,0x6E))
    add_code(doc, content)
    spacer(doc)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 3. ATTACK TIMELINE
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "3. Attack Timeline (Reconstructed)", 1, (0x1F,0x3A,0x6E))
add_para(doc,
    "Timeline dưới đây được dựng lại hoàn toàn từ Sysmon logs trên ELK. "
    "Tất cả timestamp là UTC+7 (ICT).")

spacer(doc)
timeline = [
    ("21:44:05","T+0s",  "Initial Access","Ansible controller SSH vào DESKTOP-L7FCMBQ, drop attack-simulation.ps1","sshd.exe → cmd.exe"),
    ("21:44:07","T+2s",  "Execution",     "sshd.exe → cmd.exe → powershell.exe -EncodedCommand (Ansible overhead payload)","EID-4104 ScriptBlock logged"),
    ("21:44:19","T+14s", "Execution",     "Attack script started: powershell.exe -File C:\\SocLab\\attack-simulation.ps1","EID-4104: full script content captured"),
    ("21:44:22","T+17s", "Execution",     "Stage 1: -EncodedCommand 'Write-Host Payload executed...' (simulated payload drop)","EID-1: powershell.exe parent=powershell.exe"),
    ("21:44:24","T+19s", "Discovery",     "Stage 2: cmd.exe 'whoami /all'  — account enumeration","EID-1: whoami.exe, net.exe"),
    ("21:44:25","T+20s", "Discovery",     "cmd.exe 'systeminfo'  — OS/patch level recon","EID-1: systeminfo.exe"),
    ("21:44:29","T+24s", "Discovery",     "cmd.exe 'ipconfig /all', 'netstat -ano', 'tasklist /v'  — network + process recon","EID-1: ipconfig.exe, NETSTAT.EXE, tasklist.exe"),
    ("21:44:29","T+24s", "Persistence",   "Stage 3: Set-ItemProperty HKCU:\\..\\Run\\WindowsUpdateSvc = calc.exe","EID-13: TargetObject=...\\Run\\WindowsUpdateSvc"),
    ("21:44:33","T+28s", "Persistence",   "Register-ScheduledTask 'MicrosoftEdgeUpdateTaskUA' (PS AtLogon trigger)","EID-4103: CommandInvocation(Register-ScheduledTask)"),
    ("21:44:33","T+28s", "Defense Evasion","Stage 4: New-Item $env:TEMP\\svchost32.exe.log (masquerade file)","EID-11: FileCreate T1036"),
    ("21:44:33","T+28s", "Defense Evasion","wevtutil cl Application  — Application log cleared","EID-4104: wevtutil in ScriptBlock"),
    ("21:44:34","T+29s", "Credential Access","Stage 5: OpenProcess(0x1010) on lsass.exe PID=680 by powershell.exe","EID-10: src=powershell.exe, mask=0x1010, tgt=lsass.exe"),
    ("21:44:34","T+29s", "Lateral Movement","Stage 6: net use \\\\127.0.0.1\\C$ (SMB admin share probe, failed)","EID-1: net.exe"),
    ("21:44:37","T+32s", "Command & Control","Stage 7: Invoke-WebRequest http://93.184.216.34/ (beacon ×1)","EID-3: network conn (OneDrive context; C2 beacon blocked)"),
    ("21:45:08","T+63s", "Defense Evasion","Windows Defender (MsMpEng.exe) scans lsass.exe — triggered by LSASS access","EID-10: MsMpEng → lsass 0x1010 (AV response)"),
    ("21:45:23","T+78s", "Cleanup",       "Remove-ItemProperty WindowsUpdateSvc, Unregister-ScheduledTask, Remove-Item svchost32.exe.log","EID-13, EID-4103"),
    ("21:45:28","T+83s", "End",           "Attack simulation complete. All 7 stages executed.","—"),
]

add_table(doc,
    ["Time (ICT)","T+","Tactic","Activity","Evidence/EID"],
    [[r[0],r[1],r[2],r[3],r[4]] for r in timeline],
    [2.0, 1.0, 3.0, 7.5, 4.0])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 4. ELK QUERY PLAYBOOK
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "4. ELK Query Playbook — KQL Investigation Queries", 1, (0x1F,0x3A,0x6E))
add_para(doc,
    "Phần này liệt kê từng query KQL được sử dụng, mục đích, output thu được và "
    "cách đọc kết quả. Chạy trong Kibana Discover với index pattern 'winlogbeat-*'.")

sections = [
    # (title, purpose, kql, output, insight)
    ("4.1  Xác định Attack Window — Tất cả events từ host trong giờ xảy ra",
     "Câu hỏi cần trả lời: Trong khoảng thời gian alert, có bao nhiêu events? Phân bố loại event ra sao? "
     "Đây là bước đầu tiên để hiểu scale của incident.",
     'host.name: "DESKTOP-L7FCMBQ"\nAND @timestamp >= "2026-08-17T14:44:00Z"\nAND @timestamp <= "2026-08-17T15:00:00Z"',
     "Total: 6,439 events trong ~16 phút\n"
     "EID-1  (ProcessCreate) : 122 events\n"
     "EID-4104 (PS ScriptBlock) :  89 events  ← quan trọng\n"
     "EID-10 (ProcessAccess)  :1,324 events\n"
     "EID-13 (RegistryModify) :2,666 events\n"
     "EID-11 (FileCreate)     :  395 events\n"
     "EID-3  (NetworkConn)    :    7 events",
     "Ngay lập tức thấy EID-4104 có 89 events — PowerShell ScriptBlock logging bắt được execution. "
     "EID-10 có 1,324 events — rất nhiều LSASS access cần phân tích thêm."),

    ("4.2  Stage 1 — Phát hiện PowerShell Encoded Command Execution (T1059.001)",
     "Hypothesis: Attacker dùng -EncodedCommand để obfuscate payload. "
     "EID-4104 = PowerShell Script Block logging, EID-1 = Process Create với command line.",
     'event.code: "4104"\nAND message: "EncodedCommand"\nAND @timestamp >= "2026-08-17T14:44:00Z"',
     "2026-08-17T14:44:19 UTC — ScriptBlock captured:\n"
     "  [Console]::InputEncoding = ...; powershell -ExecutionPolicy Bypass\n"
     "  -File C:\\SocLab\\attack-simulation.ps1\n\n"
     "2026-08-17T14:44:21 UTC — Full script content captured:\n"
     "  # SOC Lab — Multi-Stage Attack Simulation\n"
     "  # Simulates realistic attack chain: Execution -> Discovery\n"
     "  -> Persistence -> Credential Access",
     "EID-4104 cực kỳ giá trị — capture được toàn bộ script content dù attacker dùng -EncodedCommand. "
     "PowerShell Script Block Logging là 'game over' cho obfuscation. "
     "Note: cần bật 'Script Block Logging' trong GPO để EID-4104 được ghi."),

    ("4.3  Stage 2 — Hunt Discovery Commands (T1087/T1082/T1016)",
     "Sau khi xác nhận có execution, tìm các dấu hiệu reconnaissance. "
     "Attacker luôn enumerate môi trường ngay sau khi có foothold.",
     'event.code: "1"\nAND message: ("whoami" OR "systeminfo" OR "ipconfig"\n'
     '     OR "netstat" OR "tasklist" OR "net user")\nAND @timestamp >= "2026-08-17T14:44:00Z"',
     "21:44:24 | cmd.exe /c 'whoami /all'   parent=powershell.exe\n"
     "21:44:24 | whoami.exe                  parent=cmd.exe\n"
     "21:44:25 | cmd.exe /c systeminfo       parent=powershell.exe\n"
     "21:44:25 | systeminfo.exe              parent=cmd.exe\n"
     "21:44:29 | cmd.exe /c 'ipconfig /all'  parent=powershell.exe\n"
     "21:44:29 | ipconfig.exe                parent=cmd.exe\n"
     "21:44:29 | NETSTAT.EXE -ano            parent=cmd.exe\n"
     "21:44:29 | tasklist.exe /v             parent=cmd.exe",
     "Process tree rõ ràng: powershell.exe → cmd.exe → [recon tool]. "
     "Pattern này (PS spawns cmd spawns discovery tools) là IOA mạnh — không phải behavior bình thường. "
     "Thứ tự: whoami → systeminfo → ipconfig → netstat → tasklist = classic post-exploitation recon sequence."),

    ("4.4  Stage 3 — Persistence: Registry Run Key (T1547.001)",
     "Hunt registry modification trên autorun key. "
     "EID-13 = Sysmon registry value set. Attacker dùng Run key để survive reboot.",
     'event.code: "13"\nAND message: "\\\\CurrentVersion\\\\Run\\\\"\nAND @timestamp >= "2026-08-17T14:44:00Z"',
     "2026-08-17T14:44:29 UTC — EID-13 Registry Value Set:\n"
     "  TargetObject : HKU\\S-1-5-21-...\\SOFTWARE\\Microsoft\\Windows\\\n"
     "                 CurrentVersion\\Run\\WindowsUpdateSvc\n"
     "  Details      : C:\\Windows\\System32\\calc.exe\n"
     "  Image (by)   : powershell.exe\n\n"
     "  → Persistence established. calc.exe masquerading as 'WindowsUpdateSvc'",
     "Key indicator: 1) tên key giả mạo legitimate service (WindowsUpdateSvc), "
     "2) value trỏ đến unexpected executable (calc.exe trong môi trường thật sẽ là malware). "
     "3) Created BY powershell.exe — không phải installer hợp lệ. "
     "EID-12 (create) + EID-13 (modify) là cặp đôi quan trọng cho persistence hunting."),

    ("4.5  Stage 3b — Persistence: Scheduled Task (T1053.005)",
     "Tìm dấu hiệu scheduled task creation. "
     "EID-4103 = PowerShell Module logging capture Register-ScheduledTask cmdlet.",
     'event.code: "4103"\nAND message: "Register-ScheduledTask"\nAND @timestamp >= "2026-08-17T14:44:00Z"',
     '2026-08-17T14:44:33 UTC — EID-4103 Module Logging:\n'
     '  CommandInvocation(Register-ScheduledTask): "Register-ScheduledTask"\n'
     '  ParameterBinding: name="TaskName";\n'
     '                    value="MicrosoftEdgeUpdateTaskUA"\n'
     '  ParameterBinding: name="Action";\n'
     '                    value="powershell.exe -NonInteractive -Command whoami"\n'
     '  ParameterBinding: name="Trigger"; value="AtLogOn"',
     "TaskName 'MicrosoftEdgeUpdateTaskUA' — typosquatting tên task hợp lệ của Edge. "
     "Trigger AtLogOn = persistent. Action = powershell.exe (legitimate binary dùng để evade). "
     "KQL tương đương cho Windows Event 4698: event.code: \"4698\" AND message: \"MicrosoftEdge*\""),

    ("4.6  Stage 5 — Credential Access: LSASS Memory (T1003.001)",
     "Hunt LSASS process access với access mask Mimikatz-style (0x1010). "
     "EID-10 = Sysmon ProcessAccess. Cần filter noise từ WMI và AV trước.",
     'event.code: "10"\nAND message: "lsass.exe"\nAND message: "0x1010"\nAND @timestamp >= "2026-08-17T14:44:00Z"',
     "Event 1 (MALICIOUS — 14:44:34 UTC):\n"
     "  SourceImage : C:\\Windows\\System32\\WindowsPowerShell\\v1.0\\powershell.exe\n"
     "  TargetImage : C:\\Windows\\system32\\lsass.exe\n"
     "  GrantedAccess: 0x1010   ← PROCESS_QUERY_LIMITED_INFO | VM_READ\n"
     "  → Mimikatz-style credential dumping attempt\n\n"
     "Event 2 (BENIGN — 14:45:08 UTC):\n"
     "  SourceImage : C:\\ProgramData\\Microsoft\\Windows Defender\\...\\MsMpEng.exe\n"
     "  TargetImage : lsass.exe\n"
     "  GrantedAccess: 0x1010\n"
     "  → Windows Defender responding to the LSASS access above",
     "0x1010 mask từ powershell.exe là clear Mimikatz indicator. "
     "Thú vị: chính event LSASS access đó đã trigger Defender (MsMpEng) access LSASS 34 giây sau "
     "→ alert 'meta-evidence' tốt. "
     "Noise filter quan trọng: -message: wmiprvse AND -message: MsMpEng AND -message: wininit"),

    ("4.7  Stage 4 — Defense Evasion: Log Clear (T1070.001)",
     "Detect event log tampering. "
     "EID-1102 = Security log cleared. Nhưng attack script clear Application log (script có trong EID-4104).",
     'message: "wevtutil" AND message: ("cl" OR "clear-log")\nAND @timestamp >= "2026-08-17T14:44:00Z"',
     "2026-08-17T14:44:21 UTC — EID-4104 ScriptBlock:\n"
     "  ...wevtutil cl Application...\n"
     "  Full script captured → reveals wevtutil usage\n\n"
     "Note: EID-1102 (Security log clear) requires Security log monitoring.\n"
     "Application log clear: wevtutil cl Application — không generate EID-1102.",
     "Lesson: wevtutil cl [non-Security] không generate EID-1102. "
     "Phải hunt qua EID-4104 (script content) hoặc EID-1 (wevtutil.exe process). "
     "Best practice: hunt BOTH EID-1102 AND EID-1 với Image: wevtutil.exe để cover cả 2 case."),

    ("4.8  Stage 7 — C2 Beacon Pattern Hunt (T1071.001)",
     "Detect periodic HTTP connections to external IP. "
     "EID-3 = Sysmon NetworkConnect. Hunt outbound connections từ non-browser processes.",
     'event.code: "3"\nAND NOT message: ("chrome" OR "firefox" OR "edge" OR "MicrosoftEdge")\nAND message: ":80 "\nAND @timestamp >= "2026-08-17T14:44:00Z"',
     "7 network events trong window — tất cả từ OneDriveSetup.exe/OneDrive.Sync.Service.exe\n"
     "→ background Windows services (benign)\n\n"
     "C2 beacon (Invoke-WebRequest) của attack script → bị block bởi network firewall\n"
     "→ Không có EID-3 cho 93.184.216.34 (network layer blocked)\n\n"
     "Alternative hunting method:\n"
     "  EID-4103: message: \"Invoke-WebRequest\" AND message: \"93.184.216.34\"\n"
     "  → EID-4103 captured PowerShell module call với target IP",
     "Quan trọng: C2 beacon block ở firewall level không có nghĩa là không trace được. "
     "EID-4103 (Module Logging) capture cmdlet call dù không có kết nối mạng thực sự. "
     "Trong thực tế: cần correlate EID-3 (network) + EID-22 (DNS) để phát hiện beaconing pattern "
     "— đặc biệt là regular interval với payload nhỏ."),

    ("4.9  Attack Chain Correlation — Full Pivot Query",
     "Query tổng hợp để xem toàn bộ attack activity từ 1 host trong 1 search. "
     "Đây là query dùng khi cần 'single pane of glass' view cho một incident.",
     'host.name: "DESKTOP-L7FCMBQ"\nAND @timestamp >= "2026-08-17T14:44:00Z"\nAND (\n'
     '  (event.code: "1" AND message: ("powershell" OR "whoami" OR "systeminfo"\n'
     '       OR "netstat" OR "schtasks" OR "wevtutil"))\n'
     '  OR (event.code: "13" AND message: "CurrentVersion\\\\Run")\n'
     '  OR (event.code: "10" AND message: "lsass.exe" AND message: "0x1010")\n'
     '  OR (event.code: "4104" AND message: ("EncodedCommand" OR "Invoke-WebRequest"))\n'
     '  OR (event.code: "4103" AND message: "Register-ScheduledTask")\n'
     ')',
     "Kết quả (sort by @timestamp ascending):\n"
     "  14:44:07 EID-1    powershell.exe -EncodedCommand [Ansible overhead]\n"
     "  14:44:19 EID-4104 ScriptBlock: attack-simulation.ps1 content\n"
     "  14:44:22 EID-1    powershell.exe -EncodedCommand [Stage 1]\n"
     "  14:44:24 EID-1    whoami.exe /all            [Stage 2 Discovery]\n"
     "  14:44:25 EID-1    systeminfo.exe             [Stage 2 Discovery]\n"
     "  14:44:29 EID-1    ipconfig.exe, netstat.exe  [Stage 2 Discovery]\n"
     "  14:44:29 EID-13   Registry Run\\WindowsUpdateSvc [Stage 3 Persist]\n"
     "  14:44:33 EID-4103 Register-ScheduledTask     [Stage 3b Persist]\n"
     "  14:44:34 EID-10   powershell→lsass 0x1010   [Stage 5 CredAccess]\n"
     "  → Complete attack chain visible in chronological order",
     "Đây là 'analyst's best friend query' — thay vì chạy 7 query riêng, "
     "1 query phức hợp cho full picture. "
     "Pattern OR giữa các EID và keyword tạo 'OR-of-AND' logic hiệu quả. "
     "Recommend: save query này thành 'Saved Search' trong Kibana cho reuse."),
]

for title, purpose, kql, output, insight in sections:
    add_heading(doc, title, 2, (0x2E,0x6D,0xA3))
    add_para(doc, "📋 Mục đích:", bold=True)
    add_para(doc, purpose, indent=0.5)
    spacer(doc)
    add_para(doc, "🔍 KQL Query:", bold=True)
    add_code(doc, kql)
    spacer(doc)
    add_para(doc, "📊 Output:", bold=True)
    add_code(doc, output)
    spacer(doc)
    add_para(doc, "💡 Insight:", bold=True)
    add_para(doc, insight, indent=0.5, color=(0x2C,0x5F,0x2E))
    spacer(doc, 2)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 5. IOC / IOA SUMMARY
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "5. Indicators of Compromise (IOC) & Indicators of Attack (IOA)", 1, (0x1F,0x3A,0x6E))

add_heading(doc, "5.1  Indicators of Compromise (IOC)", 2)
add_para(doc, "IOC = artifacts tĩnh mà attacker để lại (file, registry key, network endpoint).", italic=True)
ioc_rows = [
    ("Registry Key","HKCU\\...\\CurrentVersion\\Run\\WindowsUpdateSvc","Persistence","EID-13"),
    ("Registry Value","C:\\Windows\\System32\\calc.exe","Masquerade executable","EID-13"),
    ("Scheduled Task","MicrosoftEdgeUpdateTaskUA","Persistence, typosquatting","EID-4103"),
    ("File","C:\\Users\\gnid\\AppData\\Local\\Temp\\svchost32.exe.log","Masquerade filename","EID-11"),
    ("Script","C:\\SocLab\\attack-simulation.ps1","Attack payload","EID-4104"),
    ("IP","93.184.216.34:80","C2 destination (simulated)","EID-4104 (PS log)"),
    ("IP","192.168.154.164","Compromised host","All events"),
    ("User","gnid (S-1-5-21-...-1001)","Account used during attack","EID-13, EID-10"),
]
add_table(doc, ["Type","Value","Context","Source EID"], ioc_rows, [2.5, 5.5, 4.0, 2.5])

spacer(doc)
add_heading(doc, "5.2  Indicators of Attack (IOA)", 2)
add_para(doc, "IOA = behavioral patterns — không thể fake hay giả mạo, giá trị hơn IOC.", italic=True)
ioa_rows = [
    ("powershell.exe -EncodedCommand","Obfuscated execution","T1059.001","HIGH"),
    ("sshd.exe → cmd.exe → powershell.exe","Unusual parent chain","T1059.003","HIGH"),
    ("powershell.exe → cmd.exe → whoami/systeminfo/ipconfig","Post-exploit discovery sequence","T1087+T1082+T1016","HIGH"),
    ("HKCU Run key created by powershell.exe","Persistence via script","T1547.001","HIGH"),
    ("powershell.exe OpenProcess(lsass, 0x1010)","Mimikatz-style credential dump","T1003.001","CRITICAL"),
    ("MsMpEng triggered within 34s of LSASS access","AV response = attack detected","T1003.001","HIGH"),
    ("wevtutil cl Application (in PS ScriptBlock)","Anti-forensics attempt","T1070.001","MEDIUM"),
    ("Invoke-WebRequest to external IP in script","C2 communication attempt","T1071.001","HIGH"),
]
add_table(doc, ["Behavior","Description","Technique","Severity"], ioa_rows, [5.0, 4.5, 2.5, 2.0])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 6. RECOMMENDATIONS
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "6. Recommendations", 1, (0x1F,0x3A,0x6E))

recs = [
    ("Detection Improvements",
     "• Thêm detection rule: powershell.exe với flag -EncodedCommand → alert HIGH\n"
     "• Hunt weekly: EID-13 với TargetObject chứa \\\\Run\\\\ NOT được tạo bởi installer\n"
     "• Alert ngay: EID-10 với GrantedAccess=0x1010 AND SourceImage=powershell.exe\n"
     "• Enable Windows Event 4698 (Sched Task audit) để catch schtasks qua native API"),

    ("Log Coverage Gaps",
     "• C2 beacon qua HTTP không có EID-3 nếu bị block ở firewall → cần DNS logging (EID-22)\n"
     "• PowerShell ScriptBlock logging (EID-4104) cần bật riêng trong GPO\n"
     "• Module logging (EID-4103) cần bật riêng — cho phép catch Register-ScheduledTask\n"
     "• Xem xét cài Sysmon config mới nhất (sysmonconfig_Server.xml đã tốt)"),

    ("Hardening",
     "• Restrict PowerShell: LanguageMode = ConstrainedLanguage\n"
     "• AppLocker/WDAC: chặn powershell.exe trong user AppData và Temp\n"
     "• Credential Guard: bảo vệ LSASS khỏi memory access\n"
     "• Attack Surface Reduction Rules: block Office macro spawning PS, JS\n"
     "• Network: egress filtering — chặn direct outbound HTTP từ workstation"),

    ("Incident Response Process",
     "• Bước 1: Isolate host ngay khi confirm LSASS access (EID-10 từ PS)\n"
     "• Bước 2: Snapshot memory nếu có tool (FTK, WinPmem) trước khi shutdown\n"
     "• Bước 3: Collect Sysmon EVTX, Security log, PowerShell log\n"
     "• Bước 4: Check lateral movement: tìm EID-4624 type 3 từ host này tới host khác\n"
     "• Bước 5: Reset password toàn bộ account trên host đó"),
]

for title, content in recs:
    add_para(doc, f"▶  {title}", bold=True, color=(0x1F,0x3A,0x6E))
    add_code(doc, content)
    spacer(doc)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 7. MITRE ATT&CK MAPPING
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "7. MITRE ATT&CK Coverage Matrix", 1, (0x1F,0x3A,0x6E))
add_para(doc, "Full mapping của attack simulation này tới MITRE ATT&CK Enterprise v14:")
spacer(doc)

mitre_rows = [
    ("Initial Access",     "T1078",      "Valid Accounts",                "gnid local admin used","DETECTED (logon audit)"),
    ("Execution",         "T1059.001",  "PowerShell",                    "-EncodedCommand, Bypass EP","DETECTED (EID-4104, EID-1)"),
    ("Execution",         "T1204.002",  "Malicious File",                "attack-simulation.ps1","DETECTED (EID-4104 ScriptBlock)"),
    ("Persistence",       "T1547.001",  "Registry Run Keys",             "WindowsUpdateSvc","DETECTED (EID-13)"),
    ("Persistence",       "T1053.005",  "Scheduled Task",                "MicrosoftEdgeUpdateTaskUA","DETECTED (EID-4103)"),
    ("Privilege Esc",     "T1078",      "Valid Accounts",                "Admin token via gnid","—"),
    ("Defense Evasion",   "T1070.001",  "Clear Windows Event Logs",      "wevtutil cl Application","DETECTED (EID-4104)"),
    ("Defense Evasion",   "T1036",      "Masquerading",                  "svchost32.exe.log","DETECTED (EID-11)"),
    ("Defense Evasion",   "T1027",      "Obfuscated Files/Info",         "Base64 -EncodedCommand","DETECTED (EID-4104 decoded)"),
    ("Credential Access", "T1003.001",  "LSASS Memory",                  "OpenProcess 0x1010","DETECTED (EID-10)"),
    ("Discovery",         "T1087",      "Account Discovery",             "whoami /all, net user","DETECTED (EID-1)"),
    ("Discovery",         "T1082",      "System Information",            "systeminfo, wmic","DETECTED (EID-1)"),
    ("Discovery",         "T1016",      "System Network Config",         "ipconfig, netstat, arp","DETECTED (EID-1)"),
    ("Discovery",         "T1057",      "Process Discovery",             "tasklist /v","DETECTED (EID-1)"),
    ("Lateral Movement",  "T1021.002",  "SMB/Windows Admin Shares",      "net use \\\\127.0.0.1\\C$","DETECTED (EID-1)"),
    ("Command & Control", "T1071.001",  "Web Protocols",                 "Invoke-WebRequest HTTP","PARTIAL (EID-4103, no EID-3)"),
    ("Command & Control", "T1132",      "Data Encoding",                 "Base64 in -EncodedCommand","DETECTED (EID-4104)"),
]

add_table(doc,
    ["Tactic","Technique ID","Technique Name","Implementation","Detection Status"],
    mitre_rows,
    [3.0, 2.3, 4.5, 4.5, 3.5])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 8. APPENDIX
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "Appendix A — Lab Infrastructure", 1, (0x1F,0x3A,0x6E))
infra_rows = [
    ("ELK Stack (Indexer)","43.228.215.234","Elasticsearch 8.14.3, Kibana 8.14.3, Logstash 8.14.3","VPS Ubuntu 22.04"),
    ("Windows Endpoint","192.168.154.164","Windows 10 19045, Sysmon 15.21, Winlogbeat 8.19.0","DESKTOP-L7FCMBQ"),
    ("Linux Endpoint","192.168.154.166","Ubuntu 22.04, auditd 3.0.7, Auditbeat 8.14.3","linux-ep-01"),
    ("Analyst Controller","192.168.154.151","Kali Linux, Ansible 14.0.0","Kali"),
]
add_table(doc, ["Role","IP","Components","Hostname"], infra_rows, [3.5,3.5,6.0,3.0])

spacer(doc)
add_heading(doc, "Appendix B — Key Sysmon Event IDs Reference", 1, (0x1F,0x3A,0x6E))
eid_rows = [
    ("1","Process Create","ProcessCreation","Image, CommandLine, ParentImage, ProcessGuid"),
    ("3","Network Connection","NetworkConnect","DestinationIp, DestinationPort, Image"),
    ("7","Image Loaded","ImageLoad","ImageLoaded, Signed, Signature"),
    ("8","CreateRemoteThread","CreateRemoteThread","SourceImage, TargetImage"),
    ("10","Process Access","ProcessAccess","SourceImage, TargetImage, GrantedAccess"),
    ("11","File Created","FileCreate","TargetFilename, Image"),
    ("12","Registry Create/Delete","RegistryEvent","TargetObject, Image"),
    ("13","Registry Value Set","RegistryEvent","TargetObject, Details, Image"),
    ("15","File Stream Created","FileCreateStreamHash","TargetFilename (ADS)"),
    ("22","DNS Query","DnsQuery","QueryName, Image"),
    ("23","File Delete Archived","FileDelete","TargetFilename"),
    ("25","Process Tampering","ProcessTampering","Image (hollowing detect)"),
    ("4103","PS Module Logging","—","CommandInvocation, ParameterBinding"),
    ("4104","PS ScriptBlock","—","ScriptBlockText, ScriptBlockId"),
]
add_table(doc, ["EID","Event Name","Sysmon EventType","Key Fields"], eid_rows, [1.2,4.0,3.5,6.5])

spacer(doc)
add_heading(doc, "Appendix C — ELK Access Quick Reference", 1, (0x1F,0x3A,0x6E))
add_code(doc,
    "Kibana URL  : http://43.228.215.234:5601\n"
    "Username    : elastic\n"
    "Index       : winlogbeat-*  (Windows events)\n"
    "Index       : auditbeat-*   (Linux events)\n\n"
    "Quick KQL cheat sheet:\n"
    "  All Sysmon events    : winlog.channel: \"Microsoft-Windows-Sysmon/Operational\"\n"
    "  Specific EID         : event.code: \"1\"\n"
    "  Full text search     : message: \"powershell\" AND message: \"EncodedCommand\"\n"
    "  Time filter          : @timestamp >= \"2026-08-17T14:44:00Z\"\n"
    "  Exclude noise        : NOT message: \"wmiprvse\" NOT message: \"OneDrive\"\n"
    "  Process tree pivot   : message: \"{ProcessGuid-value}\"")

# footer
spacer(doc, 2)
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer.add_run(
    f"Generated: {datetime.datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')} | "
    "SOC Lab Training — Nam @ cycloneinstruments.ai | CONFIDENTIAL")
fr.font.size = Pt(9)
fr.font.color.rgb = RGBColor(0x88,0x88,0x88)

# ── save ──────────────────────────────────────────────────────────────────────
out = "/home/kali/splunk-soc-lab/reports/DFIR_Report_INC-2026-0817-001.docx"
doc.save(out)
print(f"Saved: {out}")
print(f"Size : {__import__('os').path.getsize(out) // 1024} KB")
